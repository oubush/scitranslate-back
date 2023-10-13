import io
from docx import Document
import magic


def validate_docx_file(file: io.BytesIO) -> bool:
    """
    """
    mime_doc_type = magic.from_buffer(
        file.getbuffer().tobytes()[:2048], mime=True)
    if 'wordprocessingml.document' in mime_doc_type:
        return True
    return False


def extract_text_from_docx(file: io.BytesIO) -> str:
    """
    """
    docx_file = Document(file)
    paragraph_texts = []
    for paragraph in docx_file.paragraphs:
        paragraph_texts.append(paragraph.text)
    return ' '.join(paragraph_texts)


def convert_text_to_docx(text: str) -> io.BytesIO:
    """
    """
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    file = io.BytesIO()
    document.save(file)
    return file